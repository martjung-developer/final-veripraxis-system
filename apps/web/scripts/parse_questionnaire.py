#!/usr/bin/env python3

import json
import re
import sys
from pathlib import Path

fitz = None
Document = None
_IMPORT_ERROR = None
try:
    import fitz as _fitz
    fitz = _fitz
except Exception as exc:
    _IMPORT_ERROR = exc

try:
    from docx import Document as _Document
    Document = _Document
except Exception as exc:
    _IMPORT_ERROR = exc


RED_COLOR = 0xFF0000
BLUE_COLOR = 0x0000FF
GREEN_COLOR = 0x008000
RE_QNUM = re.compile(r"^\s*(\d{1,3})[\.\)\-]\s+(.+)$")
RE_CHOICE = re.compile(r"^\s*[\(\[]?([a-dA-D])[\.\)\]]\s+(.+)$")
# Keep section matching strict to avoid false positives on choices like "C. ..."
# Supports typical exam section headers such as I., II., III., IV., ... up to XX.
RE_SECTION_ROMAN = re.compile(
    r"^\s*(I|II|III|IV|V|VI|VII|VIII|IX|X|XI|XII|XIII|XIV|XV|XVI|XVII|XVIII|XIX|XX)\.\s+(.+?)\s*$",
    re.I,
)
RE_SECTION_PART = re.compile(r"^\s*Part\s+(\d+)\s*[-:\u2013\u2014]?\s*(.*?)\s*$", re.I)
RE_ANSWER_KEY_START = re.compile(r"^\s*(answer\s*key|key\s*answers?)\b", re.I)
RE_ANSWER_KEY_LINE = re.compile(r"^\s*(\d{1,3})\s*[\.\)\-:]?\s*([A-D])\b", re.I)
RE_SKIP = re.compile(
    r"^(bachelor|mock board|direction|name:|date:|score:|prepared by|subject reviewer)\b",
    re.I,
)
RE_AY = re.compile(r"AY \d{4}-\d{4}", re.I)


def clean_choice_prefix(text: str) -> str:
    text = re.sub(r"^[\(\[]?([a-dA-D])[\.\)\]]\s*", "", text.strip())
    return re.sub(r"\s+", " ", text).strip()


def should_skip(text: str) -> bool:
    t = text.strip()
    if len(t) < 5:
        return True
    if RE_SKIP.search(t) or RE_AY.search(t):
        return True
    if re.fullmatch(r"[\W_]+", t):
        return True
    return False


def build_row(number: int, stem: str, choices: list, passage: str = ""):
    labels = ["A", "B", "C", "D"]
    valid = True
    errors = []
    if not stem.strip():
        valid = False
        errors.append("No empty stems")
    if len(choices) < 2 or len(choices) > 5:
        valid = False
        errors.append("Each question must have 2-5 choices")
    correct = [c for c in choices if c.get("is_correct")]
    if len(correct) != 1:
        valid = False
        errors.append("Exactly 1 correct answer required")
    correct_label = correct[0]["label"] if len(correct) == 1 else ""

    by_label = {c["label"].upper(): c["text"] for c in choices}
    return {
        "_rowIndex": number + 1,
        "_valid": valid,
        "_errors": errors,
        "status": "ready" if valid else "needs_review",
        "question_text": stem.strip(),
        "question_type": "multiple_choice",
        "correct_answer": correct_label,
        "option_a": by_label.get("A", ""),
        "option_b": by_label.get("B", ""),
        "option_c": by_label.get("C", ""),
        "option_d": by_label.get("D", ""),
        "explanation": "",
        "scenario": passage.strip(),
        "difficulty": "medium",
        "points": 1,
        "exam_id": "",
        "program_id": "",
        "section_title": "",
        "section_number": None,
    }


def roman_to_int(roman: str) -> int | None:
    roman = roman.upper().strip()
    values = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}
    total = 0
    prev = 0
    for ch in reversed(roman):
        val = values.get(ch)
        if val is None:
            return None
        if val < prev:
            total -= val
        else:
            total += val
            prev = val
    return total if total > 0 else None


def parse_answer_key_map(lines: list[str]) -> dict[int, str]:
    key_map: dict[int, str] = {}
    in_key = False
    for raw in lines:
        text = re.sub(r"\s+", " ", raw).strip()
        if not text:
            continue
        if RE_ANSWER_KEY_START.search(text):
            in_key = True
            continue
        if not in_key:
            continue
        m = RE_ANSWER_KEY_LINE.match(text)
        if m:
            key_map[int(m.group(1))] = m.group(2).upper()
            continue
        # stop if key block appears to end
        if len(key_map) > 0 and RE_QNUM.match(text):
            break
    return key_map


def int_color_to_rgb(color: int) -> tuple[int, int, int]:
    color = int(color) & 0xFFFFFF
    return ((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)


def color_distance(c1: tuple[int, int, int], c2: tuple[int, int, int]) -> float:
    dr = c1[0] - c2[0]
    dg = c1[1] - c2[1]
    db = c1[2] - c2[2]
    return (dr * dr + dg * dg + db * db) ** 0.5


def is_emphasis_style(font_name: str, flags: int, color: int, baseline_color: int | None) -> bool:
    font_l = (font_name or "").lower()
    is_bold = (
        bool(flags & 16)
        or "bold" in font_l
        or "black" in font_l
        or "demi" in font_l
        or "semibold" in font_l
    )
    if is_bold:
        return True

    rgb = int_color_to_rgb(color)
    # Strong common emphasis colors.
    target_rgbs = [
        int_color_to_rgb(RED_COLOR),
        int_color_to_rgb(BLUE_COLOR),
        int_color_to_rgb(GREEN_COLOR),
    ]
    close_to_target = any(color_distance(rgb, t) <= 36 for t in target_rgbs)
    if close_to_target:
        return True

    # If document baseline color is known, treat clearly different non-black colors as emphasis.
    if baseline_color is not None:
        base_rgb = int_color_to_rgb(baseline_color)
        if color_distance(rgb, base_rgb) >= 28 and rgb != (0, 0, 0):
            return True
    return False


def split_inline_choices(line: str):
    matches = list(re.finditer(r"(?<!\w)[\(\[]?([a-dA-D])[\.\)\]]\s+", line))
    if len(matches) == 1:
        m = matches[0]
        if m.start() > 0:
            left = line[: m.start()].strip()
            right = line[m.start() :].strip()
            parts = []
            if left:
                parts.append(left)
            if right:
                parts.append(right)
            return parts if parts else [line]
    if len(matches) <= 1:
        return [line]
    parts = []
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(line)
        seg = line[start:end].strip()
        if seg:
            parts.append(seg)
    return parts


def extract_docx_lines(paragraph):
    """Return logical text lines with bold-ratio metadata from one DOCX paragraph."""
    lines = []
    buf = []
    bold_chars = 0
    total_chars = 0

    def flush():
        nonlocal buf, bold_chars, total_chars
        text = "".join(buf).strip()
        if text:
            ratio = (bold_chars / total_chars) if total_chars > 0 else 0
            lines.append({"text": re.sub(r"\s+", " ", text).strip(), "bold_ratio": ratio})
        buf = []
        bold_chars = 0
        total_chars = 0

    for run in paragraph.runs:
        raw = run.text or ""
        if not raw:
            continue
        parts = re.split(r"(\r\n|\n|\r|\v)", raw)
        for part in parts:
            if part in ("\r\n", "\n", "\r", "\v"):
                flush()
                continue
            if not part:
                continue
            buf.append(part)
            ch_count = len(part)
            total_chars += ch_count
            if run.bold is True:
                bold_chars += ch_count

    flush()
    if not lines and paragraph.text.strip():
        lines.append({"text": re.sub(r"\s+", " ", paragraph.text).strip(), "bold_ratio": 0})
    return lines


def run_is_bold(run) -> bool:
    if run.bold is True:
        return True
    if run.bold is False:
        return False
    try:
        rb = run._element.rPr.b if run._element is not None and run._element.rPr is not None else None
        if rb is not None:
            val = rb.val
            if val is None:
                return True
            return str(val).lower() not in {"0", "false", "off"}
    except Exception:
        pass
    return False


def extract_docx_lines_with_bold(paragraph):
    lines = []
    text_buf = []
    bold_buf = []

    def flush():
        nonlocal text_buf, bold_buf
        text = "".join(text_buf).strip()
        if text:
            lines.append({"text": re.sub(r"\s+", " ", text).strip(), "bold_text": "".join(bold_buf)})
        text_buf = []
        bold_buf = []

    for run in paragraph.runs:
        raw = run.text or ""
        if not raw:
            continue
        parts = re.split(r"(\r\n|\n|\r|\v)", raw)
        for part in parts:
            if part in ("\r\n", "\n", "\r", "\v"):
                flush()
                continue
            if not part:
                continue
            text_buf.append(part)
            if run_is_bold(run):
                bold_buf.append(part)
    flush()
    if not lines and paragraph.text.strip():
        lines.append({"text": re.sub(r"\s+", " ", paragraph.text).strip(), "bold_text": ""})
    return lines


def resolve_numbered_stem_and_passage(intro_lines):
    """Pick question stem from numbered block intro lines and keep prior lines as passage."""
    clean = [re.sub(r"\s+", " ", ln).strip() for ln in intro_lines if ln and ln.strip()]
    if not clean:
        return "", ""

    stem_idx = -1
    for idx in range(len(clean) - 1, -1, -1):
        if clean[idx].endswith("?"):
            stem_idx = idx
            break

    if stem_idx < 0:
        for idx in range(len(clean) - 1, -1, -1):
            if len(clean[idx]) <= 180:
                stem_idx = idx
                break

    if stem_idx < 0:
        stem_idx = len(clean) - 1

    passage = "\n".join(clean[:stem_idx]).strip()
    stem = clean[stem_idx].strip()

    tail = clean[stem_idx + 1 :]
    if tail:
        stem = f"{stem} {' '.join(tail)}".strip()

    return stem, passage


def parse_docx(path: Path):
    if Document is None:
        raise RuntimeError(f"Missing DOCX dependency: {_IMPORT_ERROR}")
    doc = Document(str(path))
    paras = [p for p in doc.paragraphs if p.text and p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text and p.text.strip():
                        paras.append(p)
    has_numbers = any(RE_QNUM.match(p.text.strip()) for p in paras[:30])

    entries = []
    for p in paras:
        ppr = p._p.pPr
        num_id = None
        ilvl = None
        if ppr is not None and ppr.numPr is not None:
            if ppr.numPr.numId is not None:
                num_id = int(ppr.numPr.numId.val)
            if ppr.numPr.ilvl is not None:
                ilvl = int(ppr.numPr.ilvl.val)
        for ln in extract_docx_lines_with_bold(p):
            if ln["text"]:
                entries.append(
                    {
                        "text": ln["text"],
                        "bold_text": ln["bold_text"],
                        "num_id": num_id,
                        "ilvl": ilvl,
                    }
                )

    numbered_starts = [e for e in entries if e["num_id"] == 1 and e["ilvl"] == 0]
    if len(numbered_starts) >= 50:
        return parse_docx_numbering_entries(entries)
    return parse_docx_numbered(paras) if has_numbers else parse_docx_unnumbered(paras)


def parse_docx_numbering_entries(entries):
    q_indices = [i for i, e in enumerate(entries) if e["num_id"] == 1 and e["ilvl"] == 0 and not should_skip(e["text"])]
    questions = []
    for qi, start in enumerate(q_indices):
        end = q_indices[qi + 1] if qi + 1 < len(q_indices) else len(entries)
        block = [e for e in entries[start:end] if not should_skip(e["text"])]
        if not block:
            continue

        first_choice_at = None
        for bi, e in enumerate(block[1:], start=1):
            is_choice_marker = bool(RE_CHOICE.match(e["text"])) or ("\t" in e["text"])
            is_numbered_choice = e["num_id"] is not None and not (e["num_id"] == 1 and e["ilvl"] == 0)
            if is_choice_marker or is_numbered_choice:
                first_choice_at = bi
                break

        if first_choice_at is None:
            stem = block[0]["text"]
            passage = ""
            choice_entries = []
        else:
            intro = block[:first_choice_at]
            choice_entries = block[first_choice_at:]
            stem = intro[0]["text"]
            passage = ""
            qline = [x["text"] for x in intro if x["text"].endswith("?")]
            if qline:
                stem = qline[-1]
                pparts = []
                for x in intro:
                    if x["text"] == stem:
                        continue
                    pparts.append(x["text"])
                passage = "\n".join(pparts).strip()
            elif len(intro) > 1:
                stem = intro[-1]["text"]
                passage = "\n".join(x["text"] for x in intro[:-1]).strip()

        choices = []
        for e in choice_entries:
            line = e["text"]
            segs = split_inline_choices(line)
            handled = False
            if len(segs) > 1:
                for seg in segs:
                    cm = RE_CHOICE.match(seg)
                    btxt = e["bold_text"].strip()
                    if cm:
                        handled = True
                        ctext = clean_choice_prefix(seg)
                        is_correct = bool(btxt) and clean_choice_prefix(btxt).lower() in ctext.lower()
                        choices.append({"label": cm.group(1).upper(), "text": ctext, "is_correct": is_correct})
                    else:
                        ctext = clean_choice_prefix(seg)
                        if ctext:
                            choices.append({"label": chr(ord("A") + len(choices)), "text": ctext, "is_correct": bool(btxt)})
            if handled:
                continue
            cm = RE_CHOICE.match(line)
            if cm:
                btxt = e["bold_text"].strip()
                ctext = clean_choice_prefix(line)
                is_correct = bool(btxt) and clean_choice_prefix(btxt).lower() in ctext.lower()
                choices.append({"label": cm.group(1).upper(), "text": ctext, "is_correct": is_correct})
                continue
            if len(line) <= 220:
                choices.append(
                    {
                        "label": chr(ord("A") + len(choices)),
                        "text": clean_choice_prefix(line),
                        "is_correct": bool(e["bold_text"].strip()),
                    }
                )

        seen = set()
        for idx, c in enumerate(choices):
            lbl = c.get("label", "")
            if lbl in seen or lbl not in {"A", "B", "C", "D", "E"}:
                c["label"] = chr(ord("A") + idx)
            seen.add(c["label"])

        if not any(c["is_correct"] for c in choices):
            aota = [c for c in choices if "all of the above" in c["text"].lower()]
            if len(aota) == 1:
                aota[0]["is_correct"] = True

        questions.append({"stem": stem, "choices": choices[:5], "passage": passage})

    rows = []
    for idx, q in enumerate(questions, start=1):
        correct_ix = [i for i, c in enumerate(q["choices"]) if c["is_correct"]]
        if len(correct_ix) > 1:
            for i in correct_ix[1:]:
                q["choices"][i]["is_correct"] = False
        rows.append(build_row(idx, q["stem"], q["choices"], q["passage"]))
    return rows


def para_is_bold(p) -> bool:
    return any((run.bold is True) and run.text.strip() for run in p.runs)


def parse_docx_numbered(paras):
    questions = []
    current = None
    pending_scenario = []

    doc_lines = []
    for p in paras:
        doc_lines.extend(extract_docx_lines_with_bold(p))

    def finalize_current():
        nonlocal current
        if not current:
            return
        stem, auto_passage = resolve_numbered_stem_and_passage(current["intro_lines"])
        current["stem"] = stem
        current["passage"] = "\n".join(filter(None, [current["passage"], auto_passage])).strip()
        questions.append(current)
        current = None

    for ln in doc_lines:
        text = ln["text"]
        if should_skip(text):
            continue

        qm = RE_QNUM.match(text)
        if qm:
            finalize_current()
            current = {
                "stem": "",
                "intro_lines": [qm.group(2).strip()],
                "choices": [],
                "passage": "\n".join(pending_scenario).strip(),
            }
            pending_scenario = []
            continue

        if current is None:
            if len(text) > 100 and not RE_CHOICE.match(text):
                pending_scenario.append(text)
            continue

        segments = split_inline_choices(text)
        parsed_choice = False
        for seg in segments:
            cm = RE_CHOICE.match(seg)
            if not cm:
                continue
            parsed_choice = True
            label = cm.group(1).upper()
            ctext = clean_choice_prefix(seg)
            if not ctext:
                continue
            is_correct = bool(ln["bold_text"].strip())
            if ln["bold_text"].strip():
                is_correct = clean_choice_prefix(ln["bold_text"]).strip().lower() in ctext.lower()
            current["choices"].append({"label": label, "text": ctext, "is_correct": is_correct})

        if not parsed_choice:
            if current["choices"]:
                current["choices"][-1]["text"] = f"{current['choices'][-1]['text']} {text}".strip()
            else:
                current["intro_lines"].append(text)

    finalize_current()

    rows = []
    for idx, q in enumerate(questions, start=1):
        # if multiple were tagged correct because same line bold, keep first
        correct_ix = [i for i, c in enumerate(q["choices"]) if c["is_correct"]]
        if len(correct_ix) > 1:
            for i in correct_ix[1:]:
                q["choices"][i]["is_correct"] = False
        if len(correct_ix) == 0:
            aota_ix = [i for i, c in enumerate(q["choices"]) if "all of the above" in c["text"].lower()]
            if len(aota_ix) == 1:
                q["choices"][aota_ix[0]]["is_correct"] = True
        rows.append(build_row(idx, q["stem"], q["choices"], q["passage"]))
    return rows


def parse_docx_unnumbered(paras):
    def is_question_stem(text: str) -> bool:
        t = text.strip()
        if not t:
            return False
        if t.endswith("?"):
            return True
        if re.match(r"^(what|which|who|when|where|why|how|to whom|in the|the following)\b", t, re.I):
            return True
        return False

    def is_probable_choice_line(text: str) -> bool:
        t = text.strip()
        if not t:
            return False
        if RE_CHOICE.match(t):
            return True
        if "\t" in t:
            return True
        if t.endswith("?"):
            return False
        return len(t) <= 180

    def has_choice_marker(text: str) -> bool:
        t = text.strip()
        if RE_CHOICE.match(t):
            return True
        if "\t" in t and re.search(r"\b([a-dA-D])[\.\)]\s+", t):
            return True
        if re.search(r"\b([a-dA-D])[\.\)]\s+", t):
            return True
        return False

    lines = []
    for p in paras:
        lines.extend(extract_docx_lines_with_bold(p))

    rows = []
    qnum = 1
    i = 0
    pending_passage = []
    while i < len(lines):
        text = lines[i]["text"]
        if should_skip(text):
            i += 1
            continue

        if not is_question_stem(text):
            look = [lines[j]["text"] for j in range(i + 1, min(i + 4, len(lines)))]
            choiceish = sum(1 for x in look if has_choice_marker(x))
            if 50 <= len(text) <= 230 and choiceish >= 1:
                stem = text
                choices = []
                i += 1
                while i < len(lines):
                    ct = lines[i]["text"]
                    if should_skip(ct):
                        i += 1
                        continue
                    if is_question_stem(ct):
                        break
                    if len(choices) >= 2 and len(ct) >= 90 and not has_choice_marker(ct):
                        break
                    if not is_probable_choice_line(ct):
                        break

                    segs = split_inline_choices(ct)
                    parsed_any = False
                    if len(segs) > 1:
                        for seg in segs:
                            cm = RE_CHOICE.match(seg)
                            btxt = lines[i]["bold_text"].strip()
                            if cm:
                                parsed_any = True
                                ctext = clean_choice_prefix(seg)
                                is_correct = bool(btxt) and clean_choice_prefix(btxt).lower() in ctext.lower()
                                choices.append({"label": cm.group(1).upper(), "text": ctext, "is_correct": is_correct})
                            else:
                                # unlabeled left-side segment in a tab pair (e.g., "Herman Hollerith   c. ...")
                                ctext = clean_choice_prefix(seg)
                                if ctext:
                                    choices.append(
                                        {
                                            "label": chr(ord("A") + len(choices)),
                                            "text": ctext,
                                            "is_correct": bool(btxt) and ctext.lower() in clean_choice_prefix(btxt).lower(),
                                        }
                                    )
                    else:
                        cm = RE_CHOICE.match(ct)
                        if cm:
                            parsed_any = True
                            ctext = clean_choice_prefix(ct)
                            btxt = lines[i]["bold_text"].strip()
                            is_correct = bool(btxt) and clean_choice_prefix(btxt).lower() in ctext.lower()
                            choices.append({"label": cm.group(1).upper(), "text": ctext, "is_correct": is_correct})
                    if not parsed_any:
                        btxt = lines[i]["bold_text"].strip()
                        choices.append({"label": chr(ord("A") + len(choices)), "text": clean_choice_prefix(ct), "is_correct": bool(btxt)})
                    i += 1

                seen = set()
                for idx, c in enumerate(choices):
                    label = c.get("label", "")
                    if label in seen or label not in {"A", "B", "C", "D", "E"}:
                        c["label"] = chr(ord("A") + idx)
                    seen.add(c["label"])

                passage = "\n".join(pending_passage).strip()
                pending_passage = []
                rows.append(build_row(qnum, stem, choices, passage))
                qnum += 1
                continue

            if len(text) > 100:
                pending_passage.append(text)
            i += 1
            continue

        stem = text
        choices = []
        i += 1
        while i < len(lines):
            ct = lines[i]["text"]
            if should_skip(ct):
                i += 1
                continue
            if is_question_stem(ct):
                break
            if not is_probable_choice_line(ct):
                if choices:
                    break
                pending_passage.append(ct)
                i += 1
                continue

            segs = split_inline_choices(ct)
            parsed_any = False
            if len(segs) > 1:
                for seg in segs:
                    cm = RE_CHOICE.match(seg)
                    if not cm:
                        continue
                    parsed_any = True
                    ctext = clean_choice_prefix(seg)
                    is_correct = False
                    btxt = lines[i]["bold_text"].strip()
                    if btxt:
                        is_correct = clean_choice_prefix(btxt).lower() in ctext.lower()
                    choices.append(
                        {
                            "label": cm.group(1).upper(),
                            "text": ctext,
                            "is_correct": is_correct,
                        }
                    )
            else:
                cm = RE_CHOICE.match(ct)
                if cm:
                    parsed_any = True
                    ctext = clean_choice_prefix(ct)
                    btxt = lines[i]["bold_text"].strip()
                    is_correct = bool(btxt) and clean_choice_prefix(btxt).lower() in ctext.lower()
                    choices.append(
                        {
                            "label": cm.group(1).upper(),
                            "text": ctext,
                            "is_correct": is_correct,
                        }
                    )

            if not parsed_any:
                # Unlabeled choice fallback for docs where label is not preserved in extracted text.
                if len(choices) >= 5:
                    break
                btxt = lines[i]["bold_text"].strip()
                choices.append(
                    {
                        "label": chr(ord("A") + len(choices)),
                        "text": clean_choice_prefix(ct),
                        "is_correct": bool(btxt),
                    }
                )
            i += 1

        # Normalize labels in sequence if labels are missing or duplicated.
        seen = set()
        for idx, c in enumerate(choices):
            label = c.get("label", "")
            if label in seen or label not in {"A", "B", "C", "D", "E"}:
                c["label"] = chr(ord("A") + idx)
            seen.add(c["label"])

        if not any(c["is_correct"] for c in choices):
            aota = [c for c in choices if "all of the above" in c["text"].lower()]
            if len(aota) == 1:
                aota[0]["is_correct"] = True

        passage = "\n".join(pending_passage).strip()
        pending_passage = []
        rows.append(build_row(qnum, stem, choices, passage))
        qnum += 1

    return rows


def parse_pdf(path: Path):
    if fitz is None:
        raise RuntimeError(f"Missing PDF dependency: {_IMPORT_ERROR}")
    doc = fitz.open(str(path))
    lines = []
    flat_text_lines: list[str] = []
    color_counts: dict[int, int] = {}
    for page in doc:
        blocks = page.get_text("dict", sort=True).get("blocks", [])
        for block in blocks:
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                parts = []
                span_infos = []
                for s in spans:
                    txt = s.get("text", "")
                    if txt:
                        parts.append(txt)
                    color = int(s.get("color", 0))
                    flags = int(s.get("flags", 0))
                    font = str(s.get("font", ""))
                    if txt and txt.strip():
                        color_counts[color] = color_counts.get(color, 0) + len(txt.strip())
                    span_infos.append(
                        {
                            "text": txt,
                            "color": color,
                            "flags": flags,
                            "font": font,
                        }
                    )
                text = re.sub(r"\s+", " ", "".join(parts)).strip()
                if text:
                    flat_text_lines.append(text)
                    lines.append({"text": text, "spans": span_infos})

    baseline_color = None
    if color_counts:
        baseline_color = max(color_counts.items(), key=lambda kv: kv[1])[0]

    questions = []
    current = None
    pending_scenario = []
    current_section_title = ""
    current_section_number = None

    for ln in lines:
        text = ln["text"]
        if should_skip(text):
            continue

        part_match = RE_SECTION_PART.match(text)
        if part_match:
            current_section_number = int(part_match.group(1))
            current_section_title = part_match.group(2).strip() or f"Part {current_section_number}"
            continue

        roman_match = RE_SECTION_ROMAN.match(text)
        if roman_match:
            roman_num = roman_to_int(roman_match.group(1))
            if roman_num is not None:
                current_section_number = roman_num
            current_section_title = roman_match.group(2).strip()
            continue

        qm = RE_QNUM.match(text)
        if qm:
            if current:
                questions.append(current)
            current = {
                "number": int(qm.group(1)),
                "stem": qm.group(2).strip(),
                "choices": [],
                "passage": "\n".join(pending_scenario).strip(),
                "section_title": current_section_title,
                "section_number": current_section_number,
            }
            pending_scenario = []
            continue
        if current is None:
            if len(text) > 100 and not RE_CHOICE.match(text):
                pending_scenario.append(text)
            continue

        if len(text) > 100 and not RE_CHOICE.match(text) and not current["choices"]:
            current["passage"] = (current["passage"] + "\n" + text).strip()
            continue

        for seg in split_inline_choices(text):
            cm = RE_CHOICE.match(seg)
            if not cm:
                continue
            raw_seg = re.sub(r"\s+", " ", seg).strip()
            raw_seg_low = raw_seg.lower()
            clean_seg = clean_choice_prefix(seg)
            seg_low = clean_seg.lower()
            score = 0
            for s in ln.get("spans", []):
                st = re.sub(r"\s+", " ", str(s.get("text", ""))).strip()
                if not st:
                    continue
                st_low = st.lower()
                overlaps = (
                    (st_low in raw_seg_low)
                    or (raw_seg_low in st_low)
                    or (st_low in seg_low)
                    or (seg_low in st_low)
                )
                if not overlaps:
                    continue
                if is_emphasis_style(
                    str(s.get("font", "")),
                    int(s.get("flags", 0)),
                    int(s.get("color", 0)),
                    baseline_color,
                ):
                    score += max(1, len(st))
            current["choices"].append(
                {
                    "label": cm.group(1).upper(),
                    "text": clean_seg,
                    "is_correct": score > 0,
                    "_emphasis_score": score,
                }
            )

    if current:
        questions.append(current)

    answer_key_map = parse_answer_key_map(flat_text_lines)
    rows = []
    for idx, q in enumerate(questions, start=1):
        correct_ix = [i for i, c in enumerate(q["choices"]) if c["is_correct"]]
        if len(correct_ix) > 1:
            best_i = max(range(len(q["choices"])), key=lambda i: q["choices"][i].get("_emphasis_score", 0))
            for i, _ in enumerate(q["choices"]):
                q["choices"][i]["is_correct"] = (i == best_i)
            correct_ix = [best_i]
        if len(correct_ix) == 0:
            keyed = answer_key_map.get(q["number"])
            if keyed:
                for choice in q["choices"]:
                    if choice["label"] == keyed:
                        choice["is_correct"] = True
                        break
                correct_ix = [i for i, c in enumerate(q["choices"]) if c["is_correct"]]
        row = build_row(idx, q["stem"], q["choices"], q["passage"])
        row["section_title"] = q.get("section_title") or ""
        row["section_number"] = q.get("section_number")
        if len(correct_ix) == 0:
            row["_valid"] = False
            row["status"] = "needs_review"
            row["_errors"] = [*row["_errors"], "Correct answer not confidently detected; needs review."]
        rows.append(row)
    return rows


def validate_batch(rows):
    issues = []
    if len(rows) != 100:
        issues.append(f"Exactly 100 questions expected; found {len(rows)}.")
    for i, row in enumerate(rows, start=1):
        if not row["_valid"]:
            issues.append(f"Row {i}: {'; '.join(row['_errors'])}")
    return issues


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: parse_questionnaire.py <file-path>"}))
        sys.exit(1)

    path = Path(sys.argv[1])
    if not path.exists():
        print(json.dumps({"error": f"File not found: {path}"}))
        sys.exit(1)

    ext = path.suffix.lower()
    try:
        if ext in [".docx", ".doc"]:
            rows = parse_docx(path)
        elif ext == ".pdf":
            rows = parse_pdf(path)
        else:
            print(json.dumps({"error": f"Unsupported format: {ext}"}))
            sys.exit(1)

        issues = validate_batch(rows)
        print(json.dumps({"rows": rows, "issues": issues}))
    except Exception as exc:
        print(json.dumps({"error": str(exc)}))
        sys.exit(1)


if __name__ == "__main__":
    main()
